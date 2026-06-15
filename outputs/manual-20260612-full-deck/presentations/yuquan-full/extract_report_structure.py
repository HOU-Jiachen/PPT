import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


def clean(text):
    return re.sub(r"\s+", " ", text or "").strip()


def iter_blocks(document):
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield "paragraph", Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield "table", Table(child, document)


def paragraph_images(paragraph):
    rel_ids = paragraph._p.xpath(".//a:blip/@r:embed")
    images = []
    for rel_id in rel_ids:
        rel = paragraph.part.rels.get(rel_id)
        if rel is not None:
            images.append(
                {
                    "rel_id": rel_id,
                    "target": str(rel.target_ref),
                }
            )
    return images


def table_data(table):
    rows = []
    for row in table.rows:
        rows.append([clean(cell.text) for cell in row.cells])
    return rows


def looks_like_heading(text, style):
    if not text:
        return False
    style_l = style.lower()
    if "heading" in style_l or "标题" in style:
        return True
    return bool(
        re.match(r"^\d+(?:\.\d+){0,3}[.、]?\s*\S+", text)
        or re.match(r"^第[一二三四五六七八九十]+[章节]\s*\S*", text)
        or re.match(r"^[一二三四五六七八九十]+、\S+", text)
    )


def main():
    source = Path(sys.argv[1])
    output = Path(sys.argv[2])
    document = Document(source)
    blocks = []
    for index, (kind, block) in enumerate(iter_blocks(document), start=1):
        if kind == "paragraph":
            text = clean(block.text)
            style = block.style.name if block.style else ""
            images = paragraph_images(block)
            if text or images:
                blocks.append(
                    {
                        "index": index,
                        "kind": kind,
                        "style": style,
                        "text": text,
                        "heading_candidate": looks_like_heading(text, style),
                        "images": images,
                    }
                )
        else:
            rows = table_data(block)
            blocks.append(
                {
                    "index": index,
                    "kind": kind,
                    "rows": rows,
                    "row_count": len(rows),
                    "column_count": max((len(row) for row in rows), default=0),
                }
            )
    payload = {
        "source": str(source.resolve()),
        "paragraph_count": len(document.paragraphs),
        "table_count": len(document.tables),
        "inline_shape_count": len(document.inline_shapes),
        "blocks": blocks,
    }
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    print(
        json.dumps(
            {
                "output": str(output.resolve()),
                "paragraphs": payload["paragraph_count"],
                "tables": payload["table_count"],
                "inline_shapes": payload["inline_shape_count"],
                "blocks": len(blocks),
            },
            ensure_ascii=False,
        )
    )


if __name__ == "__main__":
    main()
